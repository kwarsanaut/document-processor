import os
import logging
from typing import Dict, Any, Optional
import pdfplumber  # Changed from PyMuPDF
import docx
from PIL import Image
import openai
import re
from datetime import datetime
import json

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Cloud-compatible document processor for Streamlit"""
    
    def __init__(self):
        self.openai_client = None
        self._setup_openai()
        
        # Indonesian patterns
        self.indonesian_patterns = {
            'ktp': r'\b\d{16}\b',
            'npwp': r'\d{2}\.\d{3}\.\d{3}\.\d{1}-\d{3}\.\d{3}',
            'phone': r'(\+62|08)\d{2,3}[-\s]?\d{3,4}[-\s]?\d{3,4}',
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'currency': r'Rp\.?\s?[\d.,]+|IDR\s?[\d.,]+',
            'date': r'\d{1,2}[/-]\d{1,2}[/-]\d{4}|\d{4}-\d{2}-\d{2}',
            'company': r'PT\.?\s+[A-Z][A-Za-z\s]+|CV\.?\s+[A-Z][A-Za-z\s]+'
        }
    
    def _setup_openai(self):
        """Setup OpenAI client"""
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            openai.api_key = api_key
            self.openai_client = openai
            logger.info("OpenAI client initialized")
    
    def set_openai_key(self, api_key: str):
        """Set OpenAI API key"""
        openai.api_key = api_key
        self.openai_client = openai
        logger.info("OpenAI API key updated")
    
    def process_document(self, file_path: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """Main document processing pipeline"""
        if options is None:
            options = {}
        
        try:
            # Extract content based on file type
            content, metadata = self._extract_content(file_path)
            
            if not content:
                return {"error": "No content extracted from document"}
            
            result = {
                "content": content,
                "metadata": metadata,
                "extracted_at": datetime.now().isoformat()
            }
            
            # AI-powered extractions
            if options.get('extract_entities', True):
                result["entities"] = self._extract_entities(content, options.get('indonesian_mode', True))
            
            if options.get('generate_summary', True):
                result["summary"] = self._generate_summary(content)
            
            if options.get('extract_tables', True):
                result["tables"] = self._extract_tables_from_content(content)
            
            # Document classification
            result["classification"] = self._classify_document(content, metadata.get('filename', ''))
            
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {"error": str(e)}
    
    def _extract_content(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract content based on file type"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        metadata = {
            "filename": os.path.basename(file_path),
            "file_type": file_ext,
            "file_size": os.path.getsize(file_path)
        }
        
        if file_ext == '.pdf':
            return self._extract_pdf_pdfplumber(file_path, metadata)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx(file_path, metadata)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self._extract_image_basic(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _extract_pdf_pdfplumber(self, file_path: str, metadata: Dict) -> tuple[str, Dict]:
        """Extract text from PDF using pdfplumber"""
        try:
            content_parts = []
            tables_found = 0
            
            with pdfplumber.open(file_path) as pdf:
                metadata.update({
                    "pages": len(pdf.pages),
                    "pdf_info": pdf.metadata or {}
                })
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    if text and text.strip():
                        content_parts.append(f"[Page {page_num + 1}]\n{text}")
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        tables_found += len(tables)
                        for i, table in enumerate(tables):
                            if table:
                                table_text = f"\n[Table {i+1} - Page {page_num + 1}]\n"
                                for row in table:
                                    if row:
                                        table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                                content_parts.append(table_text)
                
                metadata["tables_found"] = tables_found
            
            return '\n\n'.join(content_parts), metadata
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return f"Error extracting PDF: {str(e)}", metadata
    
    def _extract_docx(self, file_path: str, metadata: Dict) -> tuple[str, Dict]:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            metadata.update({
                "paragraphs": len(doc.paragraphs),
                "tables": len(tables)
            })
            
            # Add table content to text
            if tables:
                content_parts.append("\n[TABLES]\n")
                for i, table in enumerate(tables):
                    content_parts.append(f"Table {i+1}:")
                    for row in table:
                        content_parts.append(" | ".join(row))
                    content_parts.append("")
            
            return '\n\n'.join(content_parts), metadata
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return f"Error extracting DOCX: {str(e)}", metadata
    
    def _extract_image_basic(self, file_path: str, metadata: Dict) -> tuple[str, Dict]:
        """Basic image info extraction (no OCR in cloud)"""
        try:
            with Image.open(file_path) as img:
                metadata.update({
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode
                })
            
            # Return message instead of OCR
            return f"Image uploaded: {img.width}x{img.height} {img.format}\nNote: OCR not available in cloud version", metadata
            
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            return f"Error processing image: {str(e)}", metadata
    
    def _extract_entities(self, content: str, indonesian_mode: bool = True) -> Dict[str, Any]:
        """Extract entities from content"""
        entities = {
            "people": [],
            "organizations": [],
            "dates": [],
            "monetary_amounts": [],
            "locations": [],
            "contact_info": {"emails": [], "phones": []},
            "id_numbers": []
        }
        
        if not content:
            return entities
        
        # Regex-based extraction for Indonesian patterns
        if indonesian_mode:
            # KTP numbers
            ktp_matches = re.findall(self.indonesian_patterns['ktp'], content)
            entities["id_numbers"].extend(ktp_matches)
            
            # Phone numbers
            phone_matches = re.findall(self.indonesian_patterns['phone'], content)
            entities["contact_info"]["phones"].extend(phone_matches)
            
            # Email addresses
            email_matches = re.findall(self.indonesian_patterns['email'], content)
            entities["contact_info"]["emails"].extend(email_matches)
            
            # Currency amounts
            currency_matches = re.findall(self.indonesian_patterns['currency'], content)
            entities["monetary_amounts"].extend(currency_matches)
            
            # Dates
            date_matches = re.findall(self.indonesian_patterns['date'], content)
            entities["dates"].extend(date_matches)
            
            # Company names
            company_matches = re.findall(self.indonesian_patterns['company'], content)
            entities["organizations"].extend(company_matches)
        
        # OpenAI-powered entity extraction
        if self.openai_client:
            try:
                ai_entities = self._extract_entities_with_ai(content)
                for key, values in ai_entities.items():
                    if key in entities and isinstance(values, list):
                        entities[key].extend(values)
            except Exception as e:
                logger.warning(f"AI entity extraction failed: {e}")
        
        # Clean duplicates
        for key, value in entities.items():
            if isinstance(value, list):
                entities[key] = list(set(value))
            elif isinstance(value, dict):
                for subkey, subvalue in value.items():
                    if isinstance(subvalue, list):
                        entities[key][subkey] = list(set(subvalue))
        
        return entities
    
    def _extract_entities_with_ai(self, content: str) -> Dict[str, Any]:
        """Use OpenAI to extract entities"""
        if not self.openai_client:
            return {}
        
        prompt = f"""
        Extract entities from this Indonesian business document. Return JSON format:
        {{
            "people": ["person names"],
            "organizations": ["company/org names"], 
            "locations": ["addresses, cities"],
            "dates": ["important dates"],
            "monetary_amounts": ["money amounts"]
        }}
        
        Document:
        {content[:2000]}
        """
        
        try:
            response = self.openai_client.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.1
            )
            
            result = response.choices[0].message.content
            return json.loads(result)
            
        except Exception as e:
            logger.error(f"OpenAI entity extraction failed: {e}")
            return {}
    
    def _generate_summary(self, content: str) -> Dict[str, Any]:
        """Generate document summary"""
        if not self.openai_client or not content:
            return self._generate_simple_summary(content)
        
        prompt = f"""
        Analyze this Indonesian business document and provide:
        1. Executive summary (2-3 sentences)
        2. Key points (bullet list)
        3. Urgency level (LOW/MEDIUM/HIGH/CRITICAL)
        4. Sentiment (POSITIVE/NEUTRAL/NEGATIVE)
        
        Return JSON format:
        {{
            "executive_summary": "summary text",
            "key_points": ["point1", "point2"],
            "urgency_level": "MEDIUM",
            "sentiment": "NEUTRAL"
        }}
        
        Document:
        {content[:1500]}
        """
        
        try:
            response = self.openai_client.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400,
                temperature=0.2
            )
            
            result = response.choices[0].message.content
            return json.loads(result)
            
        except Exception as e:
            logger.error(f"AI summary generation failed: {e}")
            return self._generate_simple_summary(content)
    
    def _generate_simple_summary(self, content: str) -> Dict[str, Any]:
        """Generate simple summary without AI"""
        if not content:
            return {
                "executive_summary": "No content available for summary",
                "key_points": [],
                "urgency_level": "LOW",
                "sentiment": "NEUTRAL"
            }
        
        # Simple heuristics
        word_count = len(content.split())
        sentence_count = len([s for s in content.split('.') if s.strip()])
        
        # Detect urgency keywords
        urgency_words = ['urgent', 'penting', 'segera', 'deadline', 'asap']
        urgency_level = "HIGH" if any(word in content.lower() for word in urgency_words) else "LOW"
        
        # Detect sentiment
        positive_words = ['good', 'baik', 'sukses', 'berhasil', 'positif']
        negative_words = ['bad', 'buruk', 'gagal', 'masalah', 'negatif']
        
        pos_count = sum(1 for word in positive_words if word in content.lower())
        neg_count = sum(1 for word in negative_words if word in content.lower())
        
        if pos_count > neg_count:
            sentiment = "POSITIVE"
        elif neg_count > pos_count:
            sentiment = "NEGATIVE"
        else:
            sentiment = "NEUTRAL"
        
        return {
            "executive_summary": f"Document contains {word_count} words across {sentence_count} sentences.",
            "key_points": [
                f"Document length: {word_count} words",
                f"Estimated reading time: {word_count // 200} minutes"
            ],
            "urgency_level": urgency_level,
            "sentiment": sentiment
        }
    
    def _extract_tables_from_content(self, content: str) -> list:
        """Extract table-like structures from content"""
        tables = []
        
        # Look for table patterns in content
        lines = content.split('\n')
        current_table = []
        
        for line in lines:
            # Check if line looks like table row (contains |, multiple spaces, or tabs)
            if '|' in line or '\t' in line or '  ' in line:
                # Split by common separators
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                elif '\t' in line:
                    cells = [cell.strip() for cell in line.split('\t') if cell.strip()]
                else:
                    cells = [cell.strip() for cell in re.split(r'\s{2,}', line) if cell.strip()]
                
                if len(cells) > 1:
                    current_table.append(cells)
            else:
                if current_table and len(current_table) > 1:
                    tables.append(current_table)
                current_table = []
        
        # Don't forget the last table
        if current_table and len(current_table) > 1:
            tables.append(current_table)
        
        return tables
    
    def _classify_document(self, content: str, filename: str) -> Dict[str, Any]:
        """Classify document type"""
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # Document type classification
        classifications = {
            "invoice": ["invoice", "faktur", "tagihan", "bill", "pembayaran"],
            "contract": ["contract", "kontrak", "perjanjian", "agreement"],
            "receipt": ["receipt", "kwitansi", "bukti", "struk"],
            "letter": ["surat", "letter", "memo"],
            "report": ["report", "laporan", "analisis"],
            "certificate": ["certificate", "sertifikat", "ijazah"],
            "id_document": ["ktp", "sim", "passport", "identitas"]
        }
        
        scores = {}
        for doc_type, keywords in classifications.items():
            score = 0
            for keyword in keywords:
                score += content_lower.count(keyword)
                score += filename_lower.count(keyword) * 2  # Filename has higher weight
            scores[doc_type] = score
        
        # Find best match
        if scores:
            best_type = max(scores, key=scores.get)
            confidence = min(scores[best_type] / 10, 1.0)  # Normalize confidence
        else:
            best_type = "unknown"
            confidence = 0.0
        
        return {
            "category": best_type,
            "confidence": confidence,
            "subcategory": self._get_subcategory(best_type, content)
        }
    
    def _get_subcategory(self, category: str, content: str) -> str:
        """Get more specific document subcategory"""
        content_lower = content.lower()
        
        if category == "invoice":
            if "pajak" in content_lower or "tax" in content_lower:
                return "tax_invoice"
            elif "pro forma" in content_lower:
                return "proforma_invoice"
            else:
                return "standard_invoice"
        
        elif category == "contract":
            if "kerja" in content_lower or "employment" in content_lower:
                return "employment_contract"
            elif "sewa" in content_lower or "rental" in content_lower:
                return "rental_agreement"
            else:
                return "service_contract"
        
        elif category == "letter":
            if "resign" in content_lower or "pengunduran" in content_lower:
                return "resignation_letter"
            elif "lamaran" in content_lower or "application" in content_lower:
                return "application_letter"
            else:
                return "business_letter"
        
        return "general"
